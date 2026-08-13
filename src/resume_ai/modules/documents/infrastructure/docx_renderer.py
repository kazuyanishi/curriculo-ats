from io import BytesIO

from docx import Document

from resume_ai.modules.candidate.domain.entities import Candidate, Experience


def _date(value: object) -> str:
    return str(value) if not hasattr(value, "strftime") else value.strftime("%Y-%m")


def _add_collection(document: Document, title: str, items: tuple[object, ...]) -> None:
    if not items:
        return
    document.add_heading(title, level=1)
    for item in items:
        text = item.name
        if getattr(item, "level", None) is not None:
            text += f" — {item.level.value}"
        if hasattr(item, "issuer"):
            text += f" — {item.issuer}"
        document.add_paragraph(text, style="List Bullet")


def _add_experience(document: Document, experience: Experience) -> None:
    end_date = _date(experience.end_date) if experience.end_date else "Present"
    dates = f"{_date(experience.start_date)} — {end_date}"
    document.add_paragraph(f"{experience.role} — {experience.company}")
    document.add_paragraph(dates)
    for activity in experience.activities:
        document.add_paragraph(activity.description, style="List Bullet")
    for achievement in experience.achievements:
        document.add_paragraph(achievement.description, style="List Bullet")


class DocxCandidateRenderer:
    def render(self, candidate: Candidate) -> bytes:
        document = Document()
        document.add_heading(candidate.personal_info.full_name, level=0)
        contact = [
            candidate.personal_info.city,
            candidate.personal_info.state,
            candidate.personal_info.country,
            candidate.contact_info.email,
            candidate.contact_info.phone,
            candidate.professional_links.linkedin,
            candidate.professional_links.github,
            candidate.professional_links.portfolio,
        ]
        document.add_paragraph(" | ".join(item for item in contact if item))

        if candidate.experiences:
            document.add_heading("Experience", level=1)
            for experience in candidate.experiences:
                _add_experience(document, experience)
        if candidate.education:
            document.add_heading("Education", level=1)
            for education in candidate.education:
                dates = []
                if education.start_date:
                    dates.append(_date(education.start_date))
                if education.end_date:
                    dates.append(_date(education.end_date))
                suffix = f" — {' — '.join(dates)}" if dates else ""
                document.add_paragraph(
                    f"{education.course} — {education.institution} — "
                    f"{education.status.value}{suffix}"
                )
        _add_collection(document, "Skills", candidate.skills)
        _add_collection(document, "Technologies", candidate.technologies)
        _add_collection(document, "Tools", candidate.tools)
        _add_collection(document, "Languages", candidate.languages)
        if candidate.certifications:
            document.add_heading("Certifications", level=1)
            for certification in candidate.certifications:
                parts = [certification.name, certification.issuer]
                if certification.issue_date:
                    parts.append(_date(certification.issue_date))
                if certification.expiration_date:
                    parts.append(_date(certification.expiration_date))
                if certification.credential_id:
                    parts.append(certification.credential_id)
                if certification.credential_url:
                    parts.append(certification.credential_url)
                document.add_paragraph(" — ".join(parts), style="List Bullet")
        if candidate.projects:
            document.add_heading("Projects", level=1)
            for project in candidate.projects:
                parts = [project.name, project.description]
                if project.start_date:
                    parts.append(_date(project.start_date))
                if project.end_date:
                    parts.append(_date(project.end_date))
                parts.extend(project.technologies)
                if project.url:
                    parts.append(project.url)
                document.add_paragraph(" — ".join(parts))

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
