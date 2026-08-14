from dataclasses import dataclass
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from resume_ai.integrations.ai.config import AIConfig
from resume_ai.interfaces.api.app import create_app
from resume_ai.interfaces.api.routes import _json_value
from resume_ai.modules.candidate.domain.value_objects import YearMonth
from resume_ai.modules.matching.application.semantic_schemas import SemanticMatchBatch


def _candidate_payload() -> dict:
    return {
        "personal_info": {
            "full_name": "Jane Doe",
            "city": "Curitiba",
            "state": "PR",
            "country": "Brazil",
        },
        "contact_info": {"email": "jane@example.com", "phone": "+55"},
    }


class FakeOpenAIClient:
    def __init__(self, config: AIConfig) -> None:
        self.config = config

    def generate(self, *, system_prompt: str, user_prompt: str, response_model):
        if response_model.__name__ == "SemanticMatchBatch":
            return response_model(decisions=({"criterion_index": 0, "status": "not_matched"},))
        return response_model(
            criteria=[
                {
                    "category": "technology",
                    "value": "Python",
                    "evidence": "Python is required.",
                    "importance": "required",
                }
            ]
        )


class FakeHybridOpenAIClient(FakeOpenAIClient):
    def generate(self, *, system_prompt: str, user_prompt: str, response_model):
        if response_model.__name__ == "JobCriteriaInput":
            return response_model(
                criteria=[
                    {
                        "category": "skill",
                        "value": "Infrastructure and networks",
                        "evidence": "Infrastructure and networks are required.",
                        "importance": "required",
                    }
                ]
            )
        return SemanticMatchBatch(
            decisions=(
                {
                    "criterion_index": 0,
                    "status": "matched",
                    "evidence_paths": ("experiences[0].activities[0].description",),
                },
            )
        )


def test_health_does_not_require_ai(monkeypatch) -> None:
    app = create_app()
    client = TestClient(app)

    response = client.get("/api/v1/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_json_value_serializes_year_month_before_generic_dataclass() -> None:
    @dataclass(frozen=True)
    class Example:
        month: YearMonth

    assert _json_value(YearMonth("2025-11")) == "2025-11"
    assert _json_value(Example(YearMonth("2025-11"))) == {"month": "2025-11"}


def test_analyze_serializes_year_months_in_optimized_candidate(monkeypatch) -> None:
    monkeypatch.setattr("resume_ai.bootstrap.OpenAIStructuredAIClient", FakeOpenAIClient)
    candidate = _candidate_payload()
    candidate.update(
        {
            "experiences": [
                {
                    "company": "Example Systems",
                    "role": "Backend Developer",
                    "start_date": "2025-11",
                    "end_date": "2026-01",
                }
            ],
            "education": [
                {
                    "institution": "University",
                    "course": "Computer Science",
                    "status": "completed",
                    "start_date": "2024-01",
                    "end_date": "2025-01",
                }
            ],
            "projects": [
                {
                    "name": "Resume App",
                    "description": "A resume application",
                    "start_date": "2024-02",
                    "end_date": "2024-03",
                }
            ],
        }
    )

    response = TestClient(create_app(AIConfig("key", "model"))).post(
        "/api/v1/analyze",
        json={"candidate": candidate, "job": {"description": "Python is required."}},
    )

    assert response.status_code == 200
    optimized = response.json()["optimized_candidate"]
    assert optimized["experiences"][0]["start_date"] == "2025-11"
    assert optimized["experiences"][0]["end_date"] == "2026-01"
    assert optimized["education"][0]["start_date"] == "2024-01"
    assert optimized["projects"][0]["end_date"] == "2024-03"


def test_analyze_uses_pipeline_and_returns_structured_response(monkeypatch) -> None:
    monkeypatch.setattr("resume_ai.bootstrap.OpenAIStructuredAIClient", FakeOpenAIClient)
    app = create_app(AIConfig("key", "model"))
    client = TestClient(app)

    response = client.post(
        "/api/v1/analyze",
        json={
            "candidate": _candidate_payload(),
            "job": {"description": "Python is required."},
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["criteria"]
    assert body["matching"]
    assert body["score"]
    assert "gaps" in body and "unsupported" in body["gaps"]
    assert body["optimized_candidate"]["personal_info"]["full_name"] == "Jane Doe"


def test_analyze_uses_grounded_semantic_fallback(monkeypatch) -> None:
    monkeypatch.setattr("resume_ai.bootstrap.OpenAIStructuredAIClient", FakeHybridOpenAIClient)
    candidate = _candidate_payload()
    candidate["experiences"] = [
        {
            "company": "Example Systems",
            "role": "Support Analyst",
            "start_date": "2020-01",
            "end_date": "2024-01",
            "activities": [{"description": "Configuração de redes e servidores."}],
        }
    ]

    response = TestClient(create_app(AIConfig("key", "model"))).post(
        "/api/v1/analyze",
        json={
            "candidate": candidate,
            "job": {"description": "Infrastructure and networks are required."},
        },
    )

    assert response.status_code == 200
    assert response.json()["matching"][0]["status"] == "matched"
    assert response.json()["gaps"]["gaps"] == []


def test_analyze_optimized_candidate_round_trips_to_documents(monkeypatch) -> None:
    monkeypatch.setattr("resume_ai.bootstrap.OpenAIStructuredAIClient", FakeOpenAIClient)
    client = TestClient(create_app(AIConfig("key", "model")))

    analysis_response = client.post(
        "/api/v1/analyze",
        json={
            "candidate": _candidate_payload(),
            "job": {"description": "Python is required."},
        },
    )

    assert analysis_response.status_code == 200
    optimized_candidate = analysis_response.json()["optimized_candidate"]

    docx_response = client.post("/api/v1/documents/docx", json=optimized_candidate)
    pdf_response = client.post("/api/v1/documents/pdf", json=optimized_candidate)

    assert docx_response.status_code == 200
    assert docx_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert Document(BytesIO(docx_response.content)).paragraphs
    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"].startswith("application/pdf")
    assert pdf_response.content.startswith(b"%PDF-")


def test_analyze_truth_gate_returns_422(monkeypatch) -> None:
    class HallucinatingClient(FakeOpenAIClient):
        def generate(self, *, system_prompt: str, user_prompt: str, response_model):
            return response_model(
                criteria=[
                    {
                        "category": "technology",
                        "value": "Python",
                        "evidence": "Invented.",
                    }
                ]
            )

    monkeypatch.setattr("resume_ai.bootstrap.OpenAIStructuredAIClient", HallucinatingClient)
    response = TestClient(create_app(AIConfig("key", "model"))).post(
        "/api/v1/analyze",
        json={"candidate": _candidate_payload(), "job": {"description": "Python is required."}},
    )

    assert response.status_code == 422


def test_invalid_candidate_and_job_return_422() -> None:
    client = TestClient(create_app())

    assert (
        client.post(
            "/api/v1/analyze",
            json={"candidate": {}, "job": {"description": "x"}},
        ).status_code
        == 422
    )
    assert (
        client.post(
            "/api/v1/analyze",
            json={"candidate": _candidate_payload(), "job": {"description": ""}},
        ).status_code
        == 422
    )


@pytest.mark.parametrize(
    ("path", "media_type", "signature"),
    [
        (
            "/api/v1/documents/docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            None,
        ),
        ("/api/v1/documents/pdf", "application/pdf", b"%PDF-"),
    ],
)
def test_document_endpoints_return_downloads_without_ai(
    path: str,
    media_type: str,
    signature: bytes | None,
) -> None:
    response = TestClient(create_app()).post(path, json=_candidate_payload())

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(media_type)
    expected_disposition = f'attachment; filename="resume.{path.rsplit("/", 1)[-1]}"'
    assert response.headers["content-disposition"] == expected_disposition
    assert response.content
    if signature:
        assert response.content.startswith(signature)
    else:
        assert Document(BytesIO(response.content)).paragraphs
