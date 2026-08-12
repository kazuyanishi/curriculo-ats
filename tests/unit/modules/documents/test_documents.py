from datetime import date
from io import BytesIO
from typing import get_type_hints

from docx import Document

from resume_ai.bootstrap import build_generate_candidate_documents
from resume_ai.modules.candidate.domain.entities import (
    Achievement,
    Activity,
    Candidate,
    Certification,
    ContactInfo,
    Education,
    EducationStatus,
    Experience,
    Language,
    LanguageLevel,
    PersonalInfo,
    ProfessionalLinks,
    ProficiencyLevel,
    Project,
    Skill,
    Technology,
    Tool,
)
from resume_ai.modules.documents.application.services import (
    GenerateCandidateDocuments,
    GeneratedCandidateDocuments,
    GeneratedDocument,
)
from resume_ai.modules.documents.infrastructure.docx_renderer import DocxCandidateRenderer
from resume_ai.modules.documents.infrastructure.pdf_renderer import PdfCandidateRenderer


def _candidate(*, complete: bool = True) -> Candidate:
    if not complete:
        return Candidate(
            personal_info=PersonalInfo("Jane Doe", "Curitiba", "PR", "Brazil"),
            contact_info=ContactInfo("jane@example.com", "+55"),
        )
    return Candidate(
        personal_info=PersonalInfo("Jane Doe", "Curitiba", "PR", "Brazil"),
        contact_info=ContactInfo("jane@example.com", "+55"),
        professional_links=ProfessionalLinks(linkedin="linkedin.com/jane"),
        experiences=(
            Experience(
                "Example Corp",
                "Backend Developer",
                date(2020, 1, 1),
                activities=(Activity("Developed APIs"),),
                achievements=(Achievement("Improved reliability"),),
            ),
        ),
        education=(Education("Example University", "Computer Science", EducationStatus.COMPLETED),),
        skills=(Skill("Communication", ProficiencyLevel.ADVANCED),),
        technologies=(Technology("Python"),),
        tools=(Tool("Docker"),),
        languages=(Language("English", LanguageLevel.FLUENT),),
        certifications=(Certification("AWS", "Amazon"),),
        projects=(Project("Resume API", "ATS project", technologies=("Python",)),),
    )


def test_application_preserves_renderer_bytes_and_metadata() -> None:
    candidate = _candidate(complete=False)

    class FakeRenderer:
        def __init__(self, content: bytes) -> None:
            self.content = content
            self.received = None
            self.calls = 0

        def render(self, received: Candidate) -> bytes:
            self.received = received
            self.calls += 1
            return self.content

    docx = FakeRenderer(b"docx")
    pdf = FakeRenderer(b"pdf")
    result = GenerateCandidateDocuments(docx, pdf).execute(candidate)  # type: ignore[arg-type]

    assert isinstance(result, GeneratedCandidateDocuments)
    assert result.docx == GeneratedDocument(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        b"docx",
    )
    assert result.pdf == GeneratedDocument("resume.pdf", "application/pdf", b"pdf")
    assert docx.received is pdf.received is candidate
    assert docx.calls == pdf.calls == 1


def test_docx_renderer_contains_candidate_facts_and_omits_missing_sections() -> None:
    content = DocxCandidateRenderer().render(_candidate())
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for value in (
        "Jane Doe", "jane@example.com", "Backend Developer", "Developed APIs",
        "Improved reliability", "Computer Science", "Communication", "Python",
        "Docker", "English", "AWS", "Resume API", "ATS project",
    ):
        assert value in text
    assert "Kubernetes" not in text
    minimal_document = Document(
        BytesIO(DocxCandidateRenderer().render(_candidate(complete=False)))
    )
    minimal_text = "\n".join(paragraph.text for paragraph in minimal_document.paragraphs)
    assert "Experience" not in minimal_text
    assert "Skills" not in minimal_text


def test_pdf_renderer_returns_valid_pdf_for_complete_and_minimal_candidates() -> None:
    renderer = PdfCandidateRenderer()

    assert renderer.render(_candidate()).startswith(b"%PDF-")
    assert renderer.render(_candidate(complete=False)).startswith(b"%PDF-")


def test_bootstrap_builds_concrete_document_service() -> None:
    service = build_generate_candidate_documents()

    assert isinstance(service, GenerateCandidateDocuments)
    result = service.execute(_candidate(complete=False))
    assert result.docx.content
    assert result.pdf.content.startswith(b"%PDF-")


def test_document_public_type_hints() -> None:
    hints = get_type_hints(GenerateCandidateDocuments.execute)
    assert hints["candidate"] is Candidate
    assert hints["return"] is GeneratedCandidateDocuments
