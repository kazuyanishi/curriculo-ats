from io import BytesIO

import pytest
from docx import Document
from reportlab.pdfgen.canvas import Canvas

from resume_ai.modules.candidate.application.exceptions import ResumeTextExtractionError
from resume_ai.modules.candidate.application.ports import ResumeTextExtractor
from resume_ai.modules.candidate.infrastructure.docx_text_extractor import DocxResumeTextExtractor
from resume_ai.modules.candidate.infrastructure.pdf_text_extractor import PdfResumeTextExtractor


def _docx_bytes(with_table: bool = False) -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("Jane Doe")
    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "Advanced"
        table.cell(1, 0).text = "FastAPI"
        table.cell(1, 1).text = "Intermediate"
    else:
        document.add_paragraph("Backend Developer")
        document.add_paragraph("Python")
        document.add_paragraph("FastAPI")
    document.save(output)
    return output.getvalue()


def _pdf_bytes(pages: list[list[str]]) -> bytes:
    output = BytesIO()
    canvas = Canvas(output)
    for page_lines in pages:
        for index, line in enumerate(page_lines):
            canvas.drawString(72, 750 - index * 24, line)
        canvas.showPage()
    canvas.save()
    return output.getvalue()


def _extract(extractor: ResumeTextExtractor, content: bytes) -> str:
    return extractor.extract(content)


def test_docx_extracts_paragraphs_from_memory() -> None:
    text = _extract(DocxResumeTextExtractor(), _docx_bytes())
    assert all(term in text for term in ("Jane Doe", "Backend Developer", "Python", "FastAPI"))


def test_docx_extracts_tables_in_document_order() -> None:
    text = _extract(DocxResumeTextExtractor(), _docx_bytes(with_table=True))
    assert all(
        term in text for term in ("Jane Doe", "Python", "Advanced", "FastAPI", "Intermediate")
    )
    assert text.index("Jane Doe") < text.index("Python") < text.index("FastAPI")


@pytest.mark.parametrize("content", [b"not-a-docx", b""])
def test_docx_invalid_or_empty_bytes_raise_stable_error(content: bytes) -> None:
    with pytest.raises(ResumeTextExtractionError):
        DocxResumeTextExtractor().extract(content)


def test_pdf_extracts_text_from_multiple_pages_in_order() -> None:
    text = PdfResumeTextExtractor().extract(
        _pdf_bytes([["Jane Doe", "Backend Developer"], ["Python", "FastAPI"]])
    )
    assert all(term in text for term in ("Jane Doe", "Backend Developer", "Python", "FastAPI"))
    assert text.index("Jane Doe") < text.index("Python")


def test_pdf_without_text_returns_empty_string() -> None:
    text = PdfResumeTextExtractor().extract(_pdf_bytes([[]]))
    assert text == ""


@pytest.mark.parametrize("content", [b"not-a-pdf", b""])
def test_pdf_invalid_or_empty_bytes_raise_stable_error(content: bytes) -> None:
    with pytest.raises(ResumeTextExtractionError):
        PdfResumeTextExtractor().extract(content)
